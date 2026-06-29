"""Генерация отчётов админ-панели в Excel и Word (с оформлением)."""
import datetime as dt
import io
import re

import db

# Фирменная гамма РемТехники: жёлтый / чёрный / белый
YELLOW = "FFCB05"   # титульная плашка
HEAD = "1A1A1A"     # шапки таблиц (чёрный)
BAND = "FFF6D5"     # бледно-жёлтая полоса для чередования
INK = "1A1A1A"      # текст


def _fmt_dt(s: str | None) -> str:
    if not s:
        return "—"
    m = re.match(r"(\d{4})-(\d{2})-(\d{2})[ T](\d{2}):(\d{2})", s)
    return f"{m[3]}.{m[2]}.{m[1]} {m[4]}:{m[5]}" if m else s


def _role(r: str) -> str:
    return "Администратор" if r == "admin" else "Сотрудник"


def _msg_text(content) -> str:
    if isinstance(content, str):
        return content
    if isinstance(content, list):
        return "\n".join(
            b.get("text", "") for b in content
            if isinstance(b, dict) and b.get("type") == "text"
        )
    return ""


def _gather() -> dict:
    return {
        "totals": db.admin_overview(),
        "users": db.admin_user_stats(),
        "per_day": db.messages_per_day(14),
        "activity": db.activity_log_list(limit=300),
    }


# ── Excel ─────────────────────────────────────────────────────────────────────

def build_xlsx() -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    d = _gather()
    wb = Workbook()

    thin = Side(style="thin", color="D9D9D9")
    box = Border(left=thin, right=thin, top=thin, bottom=thin)
    white_bold = Font(bold=True, color="FFFFFF")
    center = Alignment(horizontal="center", vertical="center")
    left = Alignment(horizontal="left", vertical="center")

    def title_bar(ws, text, ncols):
        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=ncols)
        c = ws.cell(1, 1, text)
        c.font = Font(bold=True, size=15, color=INK)
        c.fill = PatternFill("solid", fgColor=YELLOW)
        c.alignment = left
        ws.row_dimensions[1].height = 30
        ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=ncols)
        s = ws.cell(2, 1, f"Ремтехника · ИИ-ассистент · сформировано {dt.datetime.now():%d.%m.%Y %H:%M}")
        s.font = Font(size=10, color="7F7F7F")
        s.alignment = left

    def header(ws, row, headers):
        for i, h in enumerate(headers, 1):
            c = ws.cell(row, i, h)
            c.font = white_bold
            c.fill = PatternFill("solid", fgColor=HEAD)
            c.alignment = center
            c.border = box
        ws.row_dimensions[row].height = 20

    # ── Лист «Обзор» ──────────────────────────────────────────────────────────
    ws = wb.active
    ws.title = "Обзор"
    title_bar(ws, "Отчёт по работе ИИ-ассистента", 2)
    t = d["totals"]
    metrics = [
        ("Сотрудников", t["users"]),
        ("Диалогов", t["conversations"]),
        ("Сообщений ассистенту", t["user_messages"]),
        ("Создано документов", t["generated_files"]),
        ("Активны сегодня", t["active_today"]),
    ]
    header(ws, 4, ["Показатель", "Значение"])
    for i, (label, val) in enumerate(metrics, start=5):
        a = ws.cell(i, 1, label); a.border = box; a.alignment = left
        b = ws.cell(i, 2, val); b.border = box; b.alignment = center; b.font = Font(bold=True)
        if i % 2 == 1:
            for col in (1, 2):
                ws.cell(i, col).fill = PatternFill("solid", fgColor=BAND)
    ws.column_dimensions["A"].width = 30
    ws.column_dimensions["B"].width = 16

    # ── Лист «Сотрудники» ──────────────────────────────────────────────────────
    ws2 = wb.create_sheet("Сотрудники")
    title_bar(ws2, "Активность сотрудников", 6)
    header(ws2, 4, ["Сотрудник", "Логин", "Роль", "Диалогов", "Сообщений", "Последняя активность"])
    for idx, u in enumerate(d["users"]):
        r = 5 + idx
        vals = [u["full_name"] or u["username"], u["username"], _role(u["role"]),
                u["conversations"], u["messages"], _fmt_dt(u["last_active"])]
        for col, v in enumerate(vals, 1):
            c = ws2.cell(r, col, v)
            c.border = box
            c.alignment = center if col in (3, 4, 5) else left
            if idx % 2 == 1:
                c.fill = PatternFill("solid", fgColor=BAND)
    for col, w in zip("ABCDEF", (26, 16, 16, 11, 12, 22)):
        ws2.column_dimensions[col].width = w

    # ── Лист «По дням» ─────────────────────────────────────────────────────────
    ws3 = wb.create_sheet("По дням")
    title_bar(ws3, "Сообщения по дням (14 дней)", 2)
    header(ws3, 4, ["Дата", "Сообщений"])
    for idx, row in enumerate(d["per_day"]):
        r = 5 + idx
        a = ws3.cell(r, 1, _fmt_dt(row["day"] + " 00:00")[:10]); a.border = box; a.alignment = left
        b = ws3.cell(r, 2, row["count"]); b.border = box; b.alignment = center
        if idx % 2 == 1:
            for col in (1, 2):
                ws3.cell(r, col).fill = PatternFill("solid", fgColor=BAND)
    ws3.column_dimensions["A"].width = 16
    ws3.column_dimensions["B"].width = 14

    # ── Лист «Журнал» ──────────────────────────────────────────────────────────
    ws4 = wb.create_sheet("Журнал")
    title_bar(ws4, "Журнал активности", 4)
    header(ws4, 4, ["Время", "Сотрудник", "Действие", "Детали"])
    labels = {"login": "Вход", "register": "Регистрация", "message": "Сообщение ассистенту"}
    for idx, a in enumerate(d["activity"]):
        r = 5 + idx
        vals = [_fmt_dt(a["created_at"]), a.get("full_name") or a.get("username") or "—",
                labels.get(a["action"], a["action"]), a.get("detail") or ""]
        for col, v in enumerate(vals, 1):
            c = ws4.cell(r, col, v)
            c.border = box
            c.alignment = left
            if idx % 2 == 1:
                c.fill = PatternFill("solid", fgColor=BAND)
    for col, w in zip("ABCD", (18, 24, 22, 60)):
        ws4.column_dimensions[col].width = w

    # freeze header rows
    for sheet in wb.worksheets:
        sheet.freeze_panes = "A5"

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ── Word ──────────────────────────────────────────────────────────────────────

def build_docx() -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    d = _gather()
    navy = RGBColor(0x1A, 0x1A, 0x1A)
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    def shade(cell, color):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), color)
        tcPr.append(shd)

    def heading(text, size=15):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(size)
        r.font.color.rgb = navy
        p.space_after = Pt(6)
        return p

    def table(headers, rows, widths=None):
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Table Grid"
        for i, h in enumerate(headers):
            cell = t.rows[0].cells[i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(h)
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(10)
            shade(cell, HEAD)
        for ri, row in enumerate(rows):
            cells = t.add_row().cells
            for i, v in enumerate(row):
                cells[i].text = str(v)
                runs = cells[i].paragraphs[0].runs
                if runs:
                    runs[0].font.size = Pt(10)
                if ri % 2 == 1:
                    shade(cells[i], BAND)
        if widths:
            for row in t.rows:
                for i, w in enumerate(widths):
                    row.cells[i].width = Cm(w)
        return t

    # Заголовок — жёлтая фирменная плашка
    tbar = doc.add_table(rows=1, cols=1)
    bcell = tbar.rows[0].cells[0]
    shade(bcell, YELLOW)
    trun = bcell.paragraphs[0].add_run("Отчёт по работе ИИ-ассистента «Ремтехника»")
    trun.bold = True
    trun.font.size = Pt(16)
    trun.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    sub = doc.add_paragraph()
    sr = sub.add_run(f"Сформировано {dt.datetime.now():%d.%m.%Y %H:%M}")
    sr.font.size = Pt(10)
    sr.font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)
    doc.add_paragraph()

    # Сводка
    heading("Сводка")
    t = d["totals"]
    table(["Показатель", "Значение"], [
        ["Сотрудников", t["users"]],
        ["Диалогов", t["conversations"]],
        ["Сообщений ассистенту", t["user_messages"]],
        ["Создано документов", t["generated_files"]],
        ["Активны сегодня", t["active_today"]],
    ], widths=[8, 4])
    doc.add_paragraph()

    # Сотрудники
    heading("Активность сотрудников")
    table(
        ["Сотрудник", "Роль", "Диалогов", "Сообщений", "Последняя активность"],
        [[u["full_name"] or u["username"], _role(u["role"]), u["conversations"],
          u["messages"], _fmt_dt(u["last_active"])] for u in d["users"]],
        widths=[5, 3.5, 2.2, 2.4, 3.5],
    )
    doc.add_paragraph()

    # Журнал (последние 40)
    heading("Журнал активности (последние записи)")
    labels = {"login": "Вход", "register": "Регистрация", "message": "Сообщение ассистенту"}
    table(
        ["Время", "Сотрудник", "Действие"],
        [[_fmt_dt(a["created_at"]), a.get("full_name") or a.get("username") or "—",
          labels.get(a["action"], a["action"])] for a in d["activity"][:40]],
        widths=[3.5, 5, 5],
    )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Word: переписка одного сотрудника ─────────────────────────────────────────

def build_user_docx(user_id: int) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    user = db.get_user(user_id)
    if not user:
        raise ValueError("Сотрудник не найден")
    convs = db.admin_conversations(user_id)

    navy = RGBColor(0x1A, 0x1A, 0x1A)   # заголовки (чёрный)
    blue = RGBColor(0x1A, 0x1A, 0x1A)   # метка «Сотрудник» (чёрный)
    green = RGBColor(0x8A, 0x6D, 0x00)  # метка «Ассистент» (тёмно-жёлтый)
    grey = RGBColor(0x7F, 0x7F, 0x7F)

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    def shade(cell, color):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), color)
        tcPr.append(shd)

    # Заголовок — жёлтая фирменная плашка
    tbar = doc.add_table(rows=1, cols=1)
    bcell = tbar.rows[0].cells[0]
    shade(bcell, YELLOW)
    trun = bcell.paragraphs[0].add_run(
        f"Переписка сотрудника: {user.get('full_name') or user['username']}")
    trun.bold = True
    trun.font.size = Pt(16)
    trun.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    meta = doc.add_paragraph()
    mr = meta.add_run(
        f"Логин: {user['username']} · роль: {_role(user['role'])} · "
        f"диалогов: {len(convs)} · сформировано {dt.datetime.now():%d.%m.%Y %H:%M}"
    )
    mr.font.size = Pt(10)
    mr.font.color.rgb = grey

    if not convs:
        doc.add_paragraph("У сотрудника пока нет диалогов.")
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    for conv in convs:
        doc.add_paragraph()
        # Заголовок диалога
        ch = doc.add_paragraph()
        cr = ch.add_run(conv["title"])
        cr.bold = True
        cr.font.size = Pt(14)
        cr.font.color.rgb = navy
        sub = doc.add_paragraph()
        sr = sub.add_run(f"{conv['messages']} сообщений · обновлён {_fmt_dt(conv['updated_at'])}")
        sr.font.size = Pt(9)
        sr.font.color.rgb = grey

        messages = db.load_history(conv["id"], limit=1000)
        for m in messages:
            text = _msg_text(m["content"]).strip()
            if not text:
                continue
            is_user = m["role"] == "user"
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            label = p.add_run(("Сотрудник: " if is_user else "Ассистент: "))
            label.bold = True
            label.font.size = Pt(11)
            label.font.color.rgb = blue if is_user else green
            body = p.add_run(text)
            body.font.size = Pt(11)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

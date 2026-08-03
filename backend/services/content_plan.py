"""TASK-0901 (#49, EPIC-09) — контент-план с учётом сезонности и продуктовой линейки.

Темы и сроки придумывает модель (опираясь на базу знаний о линейке техники, акциях
и сезонности), здесь — только оформление результата в фирменном стиле: Word (.docx)
и Excel (.xlsx).

Правило задачи: без домысливания фактов о продуктах. Всё, чего нет в БЗ/запросе,
модель кладёт в `gaps` — план печатает это отдельным разделом «Требует уточнения».
"""
from __future__ import annotations

import datetime as dt
import io

from services.docx_style import BAND, DARK, HAIRLINE, YELLOW, shade

COLUMNS = ("Дата / период", "Тема", "Формат", "Продукт / линейка", "Канал")


def _items(data: dict) -> list[dict]:
    """Нормализует позиции плана (модель может прислать неполные записи)."""
    out = []
    for it in data.get("items") or []:
        if not isinstance(it, dict):
            continue
        out.append({
            "date": str(it.get("date") or "—"),
            "topic": str(it.get("topic") or ""),
            "format": str(it.get("format") or ""),
            "product": str(it.get("product") or ""),
            "channel": str(it.get("channel") or ""),
            "note": str(it.get("note") or ""),
        })
    return out


def create_content_plan(data: dict) -> bytes:
    """Контент-план в Word (фирменный стиль).

    data = {
        "title": "Контент-план на август 2026",
        "period": "01.08.2026 – 31.08.2026",
        "summary": "фокус месяца одним-двумя предложениями",
        "items": [{"date", "topic", "format", "product", "channel", "note"}],
        "seasonality": ["сезонный фактор …"],
        "gaps": ["чего нет в базе знаний — требует уточнения"],
    }
    """
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.shared import Cm, Pt, RGBColor

    ink = RGBColor(0x1A, 0x1A, 0x1A)
    grey = RGBColor(0x7F, 0x7F, 0x7F)

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    # Альбомная ориентация — в плане широкая таблица
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = section.right_margin = Cm(1.5)

    tbar = doc.add_table(rows=1, cols=1)
    bc = tbar.rows[0].cells[0]
    shade(bc, YELLOW)
    r = bc.paragraphs[0].add_run("КОНТЕНТ-ПЛАН")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = ink

    if data.get("title"):
        p = doc.add_paragraph()
        tr = p.add_run(str(data["title"]))
        tr.bold = True
        tr.font.size = Pt(13)
    meta_bits = []
    if data.get("period"):
        meta_bits.append(f"Период: {data['period']}")
    meta_bits.append(f"Сформирован {dt.datetime.now():%d.%m.%Y}")
    meta = doc.add_paragraph().add_run(" · ".join(meta_bits))
    meta.font.size = Pt(9)
    meta.font.color.rgb = grey

    if data.get("summary"):
        doc.add_paragraph(str(data["summary"]))

    items = _items(data)
    if items:
        t = doc.add_table(rows=1, cols=len(COLUMNS))
        t.style = "Table Grid"
        for i, head in enumerate(COLUMNS):
            cell = t.rows[0].cells[i]
            rr = cell.paragraphs[0].add_run(head)
            rr.bold = True
            rr.font.size = Pt(10)
            rr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            shade(cell, DARK)
        for idx, it in enumerate(items):
            cells = t.add_row().cells
            values = (it["date"], it["topic"], it["format"], it["product"], it["channel"])
            for cell, value in zip(cells, values):
                cell.text = value
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(10)
            if it["note"]:                       # примечание — мелким под темой
                note = cells[1].add_paragraph(it["note"])
                note.runs[0].font.size = Pt(8)
                note.runs[0].font.color.rgb = grey
            if idx % 2 == 0:
                for cell in cells:
                    shade(cell, "F4F4F6")
        for row in t.rows:
            for width, cell in zip((3.2, 8.6, 3.4, 5.4, 3.6), row.cells):
                cell.width = Cm(width)

    # Сезонность — почему такие темы/сроки
    seasonality = data.get("seasonality") or []
    if seasonality:
        h = doc.add_paragraph()
        hr = h.add_run("Сезонные факторы")
        hr.bold = True
        hr.font.size = Pt(12)
        for it in seasonality:
            doc.add_paragraph(str(it), style="List Bullet")

    # Пробелы — то, чего нет в БЗ (вместо выдуманных фактов о продуктах)
    gaps = data.get("gaps") or []
    if gaps:
        gt = doc.add_table(rows=1, cols=1)
        gc = gt.rows[0].cells[0]
        shade(gc, BAND)
        gr = gc.paragraphs[0].add_run("Требует уточнения (нет данных в базе знаний)")
        gr.bold = True
        gr.font.size = Pt(12)
        gr.font.color.rgb = ink
        for it in gaps:
            doc.add_paragraph(str(it), style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def create_content_plan_xlsx(data: dict) -> bytes:
    """Тот же план в Excel: лист «План» (таблица с примечаниями) + лист «Контекст»
    (сезонность и пункты для уточнения). Фирменное оформление."""
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    from openpyxl.utils import get_column_letter

    thin = Side(style="thin", color=HAIRLINE)
    box = Border(left=thin, right=thin, top=thin, bottom=thin)
    left = Alignment(horizontal="left", vertical="center", wrap_text=True)
    center = Alignment(horizontal="center", vertical="center")

    wb = Workbook()
    ws = wb.active
    ws.title = "План"

    headers = ["№", *COLUMNS, "Примечание"]
    widths = (5, 18, 46, 16, 26, 18, 38)

    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(headers))
    c = ws.cell(1, 1, str(data.get("title") or "Контент-план").upper())
    c.font = Font(bold=True, size=14, color=DARK)
    c.fill = PatternFill("solid", fgColor=YELLOW)
    c.alignment = Alignment(horizontal="left", vertical="center")
    ws.row_dimensions[1].height = 26

    row = 2
    if data.get("period"):
        ws.cell(row, 1, f"Период: {data['period']}").font = Font(size=10)
        row += 1
    if data.get("summary"):
        ws.cell(row, 1, str(data["summary"])).font = Font(size=10, italic=True)
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=len(headers))
        row += 1
    row += 1

    head_row = row
    for i, h in enumerate(headers, 1):
        cell = ws.cell(head_row, i, h)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor=DARK)
        cell.alignment = center
        cell.border = box
        ws.column_dimensions[get_column_letter(i)].width = widths[i - 1]
    row += 1

    for idx, it in enumerate(_items(data), 1):
        vals = [idx, it["date"], it["topic"], it["format"], it["product"], it["channel"], it["note"]]
        for i, v in enumerate(vals, 1):
            cell = ws.cell(row, i, v)
            cell.border = box
            cell.alignment = center if i == 1 else left
            if idx % 2 == 1:
                cell.fill = PatternFill("solid", fgColor="F4F4F6")
        row += 1
    ws.freeze_panes = ws.cell(head_row + 1, 1)   # шапка не уезжает при прокрутке

    ws2 = wb.create_sheet("Контекст")
    ws2.column_dimensions["A"].width = 30
    ws2.column_dimensions["B"].width = 96
    r2 = 1
    for label, items in (("Сезонные факторы", data.get("seasonality") or []),
                         ("Требует уточнения", data.get("gaps") or [])):
        if not items:
            continue
        cell = ws2.cell(r2, 1, label)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor=DARK)
        cell.alignment = left
        ws2.merge_cells(start_row=r2, start_column=1, end_row=r2, end_column=2)
        r2 += 1
        for it in items:
            ws2.cell(r2, 1, "•").alignment = center
            ws2.cell(r2, 2, str(it)).alignment = left
            r2 += 1
        r2 += 1

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()

"""TASK-0904 (#48, EPIC-09) — аналитический отчёт по рынку и конкурентам.

Сбор данных делает модель (веб-поиск + read_url в SSRF-контуре), здесь — только
оформление результата в фирменном стиле «Ремтехники»: Word (.docx) и Excel (.xlsx).

Важное правило задачи: без домысливания. Всё, чего не нашлось в источниках, модель
кладёт в `gaps` — и отчёт печатает это отдельным разделом «Пробелы данных», а не
заполняет пустые места догадками. Каждая позиция несёт ссылку на источник.
"""
from __future__ import annotations

import datetime as dt
import io

from services.docx_style import BAND, DARK, HAIRLINE, YELLOW, shade

TITLE_DEFAULT = "Анализ рынка и конкурентов"


def _rows(data: dict) -> list[dict]:
    """Нормализует список конкурентов/позиций (терпим к неполным данным модели)."""
    out = []
    for c in data.get("competitors") or []:
        if not isinstance(c, dict):
            continue
        out.append({
            "name": str(c.get("name") or "—"),
            "position": str(c.get("position") or ""),
            "price_range": str(c.get("price_range") or ""),
            "note": str(c.get("note") or ""),
            "source": str(c.get("source") or ""),
        })
    return out


def create_market_report(data: dict) -> bytes:
    """Аналитический отчёт по рынку (Word, фирменный стиль).

    data = {
        "title": "Анализ рынка экскаваторов 20 т",
        "summary": "главный вывод 1–3 предложениями",
        "competitors": [{"name", "position", "price_range", "note", "source"}],
        "trends": ["тенденция …"], "observations": ["наблюдение …"],
        "recommendations": ["что делать …"],
        "gaps": ["чего не нашлось в источниках"],
        "sources": [{"title": "…", "url": "https://…"}],
    }
    """
    from docx import Document
    from docx.shared import Cm, Pt, RGBColor

    ink = RGBColor(0x1A, 0x1A, 0x1A)
    grey = RGBColor(0x7F, 0x7F, 0x7F)

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    # Фирменная жёлтая плашка-заголовок
    tbar = doc.add_table(rows=1, cols=1)
    bc = tbar.rows[0].cells[0]
    shade(bc, YELLOW)
    r = bc.paragraphs[0].add_run("АНАЛИЗ РЫНКА И КОНКУРЕНТОВ")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = ink

    if data.get("title"):
        p = doc.add_paragraph()
        tr = p.add_run(str(data["title"]))
        tr.bold = True
        tr.font.size = Pt(13)
    meta = doc.add_paragraph().add_run(f"Сформировано {dt.datetime.now():%d.%m.%Y}")
    meta.font.size = Pt(9)
    meta.font.color.rgb = grey

    if data.get("summary"):
        doc.add_paragraph(str(data["summary"]))

    # Таблица конкурентов/позиций
    rows = _rows(data)
    if rows:
        h = doc.add_paragraph()
        hr = h.add_run("Конкуренты и позиции")
        hr.bold = True
        hr.font.size = Pt(13)

        headers = ("Компания / позиция", "Специализация", "Диапазон цен", "Наблюдение")
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Table Grid"
        for i, head in enumerate(headers):
            cell = t.rows[0].cells[i]
            rr = cell.paragraphs[0].add_run(head)
            rr.bold = True
            rr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            shade(cell, DARK)
        for idx, c in enumerate(rows):
            cells = t.add_row().cells
            cells[0].text = c["name"]
            cells[1].text = c["position"]
            cells[2].text = c["price_range"] or "—"
            cells[3].text = c["note"]
            if idx % 2 == 0:                       # мягкое чередование строк
                for cell in cells:
                    shade(cell, "F4F4F6")
        for row in t.rows:                          # ширины колонок
            for width, cell in zip((5.0, 3.6, 3.4, 5.0), row.cells):
                cell.width = Cm(width)

        # Источники по позициям — отдельной строкой, чтобы ссылки не рвали таблицу
        with_src = [c for c in rows if c["source"]]
        if with_src:
            sp = doc.add_paragraph()
            spr = sp.add_run("Источники по позициям:")
            spr.bold = True
            spr.font.size = Pt(10)
            for c in with_src:
                line = doc.add_paragraph(f"{c['name']} — {c['source']}", style="List Bullet")
                line.runs[0].font.size = Pt(9)
                line.runs[0].font.color.rgb = grey

    # Текстовые разделы
    for label, items in (
        ("Тенденции рынка", data.get("trends") or []),
        ("Наблюдения", data.get("observations") or []),
        ("Рекомендации", data.get("recommendations") or []),
    ):
        if not items:
            continue
        h = doc.add_paragraph()
        hr = h.add_run(label)
        hr.bold = True
        hr.font.size = Pt(13)
        for it in items:
            doc.add_paragraph(str(it), style="List Bullet")

    # Пробелы данных — то, чего НЕ нашлось (вместо домысливания)
    gaps = data.get("gaps") or []
    if gaps:
        gt = doc.add_table(rows=1, cols=1)
        gc = gt.rows[0].cells[0]
        shade(gc, BAND)
        gr = gc.paragraphs[0].add_run("Пробелы данных (требует уточнения)")
        gr.bold = True
        gr.font.size = Pt(12)
        gr.font.color.rgb = ink
        for it in gaps:
            doc.add_paragraph(str(it), style="List Bullet")

    # Список источников
    sources = [s for s in (data.get("sources") or []) if isinstance(s, dict)]
    if sources:
        h = doc.add_paragraph()
        hr = h.add_run("Источники")
        hr.bold = True
        hr.font.size = Pt(13)
        for s in sources:
            title = str(s.get("title") or s.get("url") or "").strip()
            url = str(s.get("url") or "").strip()
            text = f"{title} — {url}" if title and url and title != url else (title or url)
            para = doc.add_paragraph(text, style="List Bullet")
            para.runs[0].font.size = Pt(9)
            para.runs[0].font.color.rgb = grey

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def create_market_report_xlsx(data: dict) -> bytes:
    """Тот же отчёт в Excel: лист «Конкуренты» (таблица) + лист «Выводы»
    (тенденции/наблюдения/рекомендации/пробелы/источники). Фирменное оформление."""
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    from openpyxl.utils import get_column_letter

    thin = Side(style="thin", color=HAIRLINE)
    box = Border(left=thin, right=thin, top=thin, bottom=thin)
    left = Alignment(horizontal="left", vertical="center", wrap_text=True)
    center = Alignment(horizontal="center", vertical="center")

    wb = Workbook()
    ws = wb.active
    ws.title = "Конкуренты"

    headers = ["№", "Компания / позиция", "Специализация", "Диапазон цен", "Наблюдение", "Источник"]
    widths = (5, 32, 24, 20, 42, 44)

    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(headers))
    c = ws.cell(1, 1, str(data.get("title") or TITLE_DEFAULT).upper())
    c.font = Font(bold=True, size=14, color=DARK)
    c.fill = PatternFill("solid", fgColor=YELLOW)
    c.alignment = Alignment(horizontal="left", vertical="center")
    ws.row_dimensions[1].height = 26

    row = 3
    if data.get("summary"):
        ws.cell(row, 1, str(data["summary"])).font = Font(size=10, italic=True)
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=len(headers))
        row += 2

    head_row = row
    for i, h in enumerate(headers, 1):
        cell = ws.cell(head_row, i, h)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor=DARK)
        cell.alignment = center
        cell.border = box
        ws.column_dimensions[get_column_letter(i)].width = widths[i - 1]
    row += 1

    for idx, comp in enumerate(_rows(data), 1):
        vals = [idx, comp["name"], comp["position"], comp["price_range"] or "—",
                comp["note"], comp["source"]]
        for i, v in enumerate(vals, 1):
            cell = ws.cell(row, i, v)
            cell.border = box
            cell.alignment = center if i == 1 else left
            if idx % 2 == 1:
                cell.fill = PatternFill("solid", fgColor="F4F4F6")
        row += 1

    # Лист выводов
    ws2 = wb.create_sheet("Выводы")
    ws2.column_dimensions["A"].width = 26
    ws2.column_dimensions["B"].width = 96
    r2 = 1
    blocks = [
        ("Тенденции рынка", data.get("trends") or []),
        ("Наблюдения", data.get("observations") or []),
        ("Рекомендации", data.get("recommendations") or []),
        ("Пробелы данных", data.get("gaps") or []),
        ("Источники", [f"{s.get('title', '')} — {s.get('url', '')}".strip(" —")
                       for s in (data.get("sources") or []) if isinstance(s, dict)]),
    ]
    for label, items in blocks:
        if not items:
            continue
        cell = ws2.cell(r2, 1, label)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor=DARK)
        cell.alignment = left
        ws2.merge_cells(start_row=r2, start_column=1, end_row=r2, end_column=2)
        r2 += 1
        for it in items:
            ws2.cell(r2, 1, "•").alignment = center
            ws2.cell(r2, 2, str(it)).alignment = left
            r2 += 1
        r2 += 1

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()

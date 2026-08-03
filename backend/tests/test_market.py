"""TASK-0904 (#48) — аналитический отчёт по рынку и конкурентам: оформление
(.docx/.xlsx), проводка инструмента, отсутствие домысливания (gaps), RBAC."""
import io

from agent.registry import role_can_use_tool
from services import docgen

FIXTURE = {
    "title": "Рынок экскаваторов 20 т, Красноярский край",
    "summary": "Предложение стабильно, цены в диапазоне 8–11 млн ₽.",
    "competitors": [
        {"name": "ООО «СпецТехТорг»", "position": "дилер SANY", "price_range": "8,5–9,8 млн ₽",
         "note": "склад в Красноярске", "source": "https://example.com/a"},
        {"name": "ООО «УралМаш-Сервис»", "position": "б/у техника", "price_range": "5–7 млн ₽",
         "note": "гарантия 6 мес.", "source": "https://example.com/b"},
    ],
    "trends": ["Рост доли китайских брендов"],
    "observations": ["Сроки поставки сократились до 2–3 недель"],
    "recommendations": ["Держать 2 машины на складе"],
    "gaps": ["Не найдено: доли рынка по дилерам за 2026 г."],
    "sources": [{"title": "Отраслевой обзор", "url": "https://example.com/report"}],
}


def test_market_report_docx_has_all_sections():
    from docx import Document
    out = docgen.create_market_report(FIXTURE)
    assert out[:2] == b"PK" and len(out) > 5000
    doc = Document(io.BytesIO(out))
    text = "\n".join(p.text for p in doc.paragraphs)
    tables = " ".join(c.text for t in doc.tables for r in t.rows for c in r.cells)
    both = text + " " + tables
    assert "АНАЛИЗ РЫНКА И КОНКУРЕНТОВ" in tables          # фирменная плашка
    assert FIXTURE["title"] in text and "8–11 млн" in text  # тема и резюме
    assert "СпецТехТорг" in tables and "8,5–9,8 млн ₽" in tables
    assert "Рост доли китайских брендов" in text           # тенденции
    assert "Держать 2 машины" in text                      # рекомендации
    assert "Пробелы данных" in both and "доли рынка" in text
    assert "https://example.com/report" in text            # источники со ссылками


def test_market_report_xlsx_two_sheets():
    from openpyxl import load_workbook
    out = docgen.create_market_report_xlsx(FIXTURE)
    wb = load_workbook(io.BytesIO(out))
    assert wb.sheetnames == ["Конкуренты", "Выводы"]
    comp = " ".join(str(c.value) for row in wb["Конкуренты"].iter_rows() for c in row if c.value)
    assert "СпецТехТорг" in comp and "https://example.com/a" in comp
    concl = " ".join(str(c.value) for row in wb["Выводы"].iter_rows() for c in row if c.value)
    assert "Пробелы данных" in concl and "Рекомендации" in concl


async def test_analyze_market_tool_both_formats(monkeypatch):
    import app.orchestrator as orch
    saved = []

    async def fake_save(self, uid, cid, name, data, kind, emit, etype):
        saved.append((name, kind, data[:2]))

    monkeypatch.setattr(orch.Orchestrator, "_save_file", fake_save)

    async def emit(_e):
        pass
    res = await orch.Orchestrator()._execute_tool(
        "analyze_market", {**FIXTURE, "filename": "Рынок", "format": "both"},
        emit, 1, None, None)
    assert [n for n, _, _ in saved] == ["Рынок.docx", "Рынок.xlsx"]
    assert all(head == b"PK" for _, _, head in saved)
    assert "2 позиций" in res and "пробелов данных: 1" in res


async def test_analyze_market_no_data_does_not_crash(monkeypatch):
    """Источник недоступен / веб-поиск ничего не дал: отчёт всё равно собирается
    (с пометкой пробелов), ход не падает и ничего не выдумывается."""
    from docx import Document

    import app.orchestrator as orch
    captured = {}

    async def fake_save(self, uid, cid, name, data, kind, emit, etype):
        captured["data"] = data

    monkeypatch.setattr(orch.Orchestrator, "_save_file", fake_save)

    async def emit(_e):
        pass
    res = await orch.Orchestrator()._execute_tool(
        "analyze_market",
        {"title": "Рынок погрузчиков", "competitors": [], "sources": [],
         "gaps": ["Источники недоступны: цены не найдены"]},
        emit, 1, None, None)
    assert "0 позиций" in res and "пробелов данных: 1" in res
    text = "\n".join(p.text for p in Document(io.BytesIO(captured["data"])).paragraphs)
    assert "Источники недоступны" in text          # пробел явно помечен
    assert "млн" not in text                        # цифры не домыслены


def test_analyze_market_rbac():
    assert role_can_use_tool("маркетинг", "analyze_market") is True
    assert role_can_use_tool("руководство", "analyze_market") is True
    assert role_can_use_tool("admin", "analyze_market") is True
    assert role_can_use_tool("продажи", "analyze_market") is False
    assert role_can_use_tool("user", "analyze_market") is False

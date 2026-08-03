"""TASK-0901 (#49) — контент-план: оформление (.docx/.xlsx), проводка инструмента,
пометка недостающих данных (без домысливания), RBAC."""
import io

from agent.registry import role_can_use_tool
from services import docgen

FIXTURE = {
    "title": "Контент-план на август 2026",
    "period": "01.08.2026 – 31.08.2026",
    "summary": "Фокус — подготовка парка к осеннему сезону.",
    "items": [
        {"date": "05.08", "topic": "ТО экскаватора перед сезоном дождей", "format": "пост",
         "product": "XCMG XE215C", "channel": "Telegram", "note": "с чек-листом"},
        {"date": "12.08", "topic": "Кейс: погрузчик на лесозаготовке", "format": "видео",
         "product": "LiuGong 856H", "channel": "VK"},
        {"date": "26.08", "topic": "Склад запчастей: наличие фильтров", "format": "сторис",
         "product": "запчасти", "channel": "Telegram"},
    ],
    "seasonality": ["Август — подготовка техники к осенне-зимнему периоду"],
    "gaps": ["Уточнить: условия акции на ТО в августе (нет в базе знаний)"],
}


def test_content_plan_docx_structure():
    from docx import Document
    out = docgen.create_content_plan(FIXTURE)
    assert out[:2] == b"PK" and len(out) > 5000
    doc = Document(io.BytesIO(out))
    text = "\n".join(p.text for p in doc.paragraphs)
    tables = " ".join(c.text for t in doc.tables for r in t.rows for c in r.cells)
    both = text + " " + tables
    assert "КОНТЕНТ-ПЛАН" in tables                      # фирменная плашка
    assert FIXTURE["title"] in text and "01.08.2026" in text
    assert "ТО экскаватора" in tables and "XCMG XE215C" in tables
    assert "видео" in tables and "Telegram" in tables     # формат и канал
    assert "Август — подготовка" in text                 # сезонность
    assert "Требует уточнения" in both and "условия акции" in both


def test_content_plan_xlsx_two_sheets():
    from openpyxl import load_workbook
    out = docgen.create_content_plan_xlsx(FIXTURE)
    wb = load_workbook(io.BytesIO(out))
    assert wb.sheetnames == ["План", "Контекст"]
    plan = " ".join(str(c.value) for row in wb["План"].iter_rows() for c in row if c.value)
    assert "Кейс: погрузчик" in plan and "LiuGong 856H" in plan and "с чек-листом" in plan
    ctx = " ".join(str(c.value) for row in wb["Контекст"].iter_rows() for c in row if c.value)
    assert "Сезонные факторы" in ctx and "Требует уточнения" in ctx


async def test_create_content_plan_tool_both_formats(monkeypatch):
    import app.orchestrator as orch
    saved = []

    async def fake_save(self, uid, cid, name, data, kind, emit, etype):
        saved.append((name, kind, data[:2]))

    monkeypatch.setattr(orch.Orchestrator, "_save_file", fake_save)

    async def emit(_e):
        pass
    res = await orch.Orchestrator()._execute_tool(
        "create_content_plan", {**FIXTURE, "filename": "План_август", "format": "both"},
        emit, 1, None, None)
    assert [n for n, _, _ in saved] == ["План_август.docx", "План_август.xlsx"]
    assert all(head == b"PK" for _, _, head in saved)
    assert "3 публикаций" in res and "уточнения пунктов: 1" in res


async def test_content_plan_marks_missing_data(monkeypatch):
    """Данных о продукте нет в БЗ: план собирается, факты НЕ выдумываются —
    недостающее уходит в раздел «Требует уточнения»."""
    from docx import Document

    import app.orchestrator as orch
    captured = {}

    async def fake_save(self, uid, cid, name, data, kind, emit, etype):
        captured["data"] = data

    monkeypatch.setattr(orch.Orchestrator, "_save_file", fake_save)

    async def emit(_e):
        pass
    res = await orch.Orchestrator()._execute_tool(
        "create_content_plan",
        {"title": "План на сентябрь",
         "items": [{"date": "02.09", "topic": "Анонс новинки", "format": "пост"}],
         "gaps": ["Уточнить: модель и цена новинки — нет в базе знаний"]},
        emit, 1, None, None)
    assert "1 публикаций" in res and "уточнения пунктов: 1" in res
    doc = Document(io.BytesIO(captured["data"]))
    text = "\n".join(p.text for p in doc.paragraphs)
    tables = " ".join(c.text for t in doc.tables for r in t.rows for c in r.cells)
    assert "Уточнить: модель и цена" in (text + " " + tables)
    assert "₽" not in text                    # цена не домыслена


def test_content_plan_rbac():
    assert role_can_use_tool("маркетинг", "create_content_plan") is True
    assert role_can_use_tool("руководство", "create_content_plan") is True
    assert role_can_use_tool("admin", "create_content_plan") is True
    assert role_can_use_tool("продажи", "create_content_plan") is False

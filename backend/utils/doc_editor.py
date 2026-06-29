"""
Hash-based Word document editor.
Uses docx_editor for stable paragraph references, falls back to python-docx.
"""
import io
import logging
import os
import tempfile

logger = logging.getLogger(__name__)


def read_doc(file_bytes: bytes) -> str:
    """
    Returns all paragraphs with hash-anchored refs.
    Format: "P1#a7b2 | Текст параграфа..."
    Tables included as "T{row}C{col}#hash | text"
    """
    try:
        return _read_with_docx_editor(file_bytes)
    except Exception as e:
        logger.warning("docx_editor read failed, fallback: %s", e)
        return _read_fallback(file_bytes)


def apply_doc_edits(file_bytes: bytes, operations: list[dict], filename: str = "document") -> tuple[bytes, str]:
    """
    Apply edits by hash-anchored refs. Operations:
      {"op": "rewrite", "ref": "P9#f3c1", "new_text": "новый текст"}
      {"op": "delete",  "ref": "P9#f3c1"}
      {"op": "insert_after", "ref": "P5#a1b2", "text": "новый абзац"}
    Returns (new_docx_bytes, diff_summary).
    """
    try:
        return _apply_with_docx_editor(file_bytes, operations, filename)
    except Exception as e:
        logger.warning("docx_editor apply failed, fallback: %s", e)
        return _apply_fallback(file_bytes, operations, filename)


# ── docx_editor implementation ─────────────────────────────────────────────

def _read_with_docx_editor(file_bytes: bytes) -> str:
    from docx import Document as PD
    from docx_editor import Document as DE
    from docx.oxml.ns import qn

    # Use python-docx to find which w:p elements are inside tables
    pd_doc = PD(io.BytesIO(file_bytes))
    table_para_set: set = set()
    for table in pd_doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    table_para_set.add(id(p._element))

    # Get all w:p elements in document order
    all_wp = list(pd_doc.element.body.iter(qn("w:p")))
    table_indices: set[int] = set()
    for i, wp in enumerate(all_wp):
        for tp in pd_doc.tables:
            for row in tp.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if p._element is wp:
                            table_indices.add(i)

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name
    try:
        de_doc = DE.open(tmp_path)
        refs = de_doc.list_paragraphs(max_chars=120)
        de_doc.close()
    finally:
        os.unlink(tmp_path)

    if not refs:
        return "Документ пустой или не содержит текстовых параграфов."

    # Filter: only top-level paragraphs (not inside tables)
    top_level = []
    table_count = len(pd_doc.tables)
    for i, ref in enumerate(refs):
        if i not in table_indices:
            top_level.append(ref)

    lines = [
        f"Документ: {len(refs)} параграфов всего, {len(top_level)} верхнеуровневых, {table_count} таблиц.",
        f"Показаны только верхнеуровневые параграфы — таблицы сохраняются автоматически при редактировании.",
        ""
    ]
    lines += top_level
    return "\n".join(lines)


def _apply_with_docx_editor(file_bytes: bytes, operations: list[dict], filename: str) -> tuple[bytes, str]:
    """
    Read paragraph list via docx_editor (for stable hashes),
    apply all changes via python-docx (rewrite/delete/insert).
    """
    import copy
    from docx import Document as PD
    from docx_editor import Document as DE
    from lxml import etree

    # ── Step 1: get paragraph refs from docx_editor ──────────────────────
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name
    try:
        de_doc = DE.open(tmp_path)
        refs_list = de_doc.list_paragraphs(max_chars=0)   # ["P1#a7b2", ...]
        de_doc.close()
    finally:
        os.unlink(tmp_path)

    # Build index map using ALL w:p elements (same as docx_editor)
    # docx_editor counts ALL w:p in the XML tree, including inside tables/headers
    # python-docx doc.paragraphs only returns top-level paragraphs — MISMATCH!
    # So we get all w:p from the XML directly.
    doc = PD(io.BytesIO(file_bytes))
    from docx.oxml.ns import qn as _qn
    all_paras = doc.element.body.iter(_qn("w:p"))
    all_paras_list = list(all_paras)

    ref_to_idx: dict[str, int] = {}
    hash_to_idx: dict[str, int] = {}
    for entry in refs_list:
        try:
            p_part, h = entry.split("#", 1)
            idx = int(p_part[1:]) - 1        # P9 → index 8 (0-based)
            ref_to_idx[entry] = idx
            if h not in hash_to_idx:          # first occurrence wins
                hash_to_idx[h] = idx
        except Exception:
            pass

    def resolve(ref: str) -> int:
        if ref in ref_to_idx:
            return ref_to_idx[ref]
        h = ref.split("#", 1)[1] if "#" in ref else ""
        return hash_to_idx.get(h, -1)

    def get_para_elem(idx: int):
        if 0 <= idx < len(all_paras_list):
            return all_paras_list[idx]
        return None

    def is_in_table(elem) -> bool:
        """Check if paragraph is inside a table cell (w:tc)."""
        parent = elem.getparent()
        while parent is not None:
            if parent.tag == _qn("w:tc"):
                return True
            parent = parent.getparent()
        return False

    def get_containing_table(elem):
        """Return the w:tbl element containing this paragraph, or None."""
        parent = elem.getparent()
        while parent is not None:
            if parent.tag == _qn("w:tbl"):
                return parent
            parent = parent.getparent()
        return None

    done: list[str] = []
    errors: list[str] = []

    rewrites = [op for op in operations if op.get("op") == "rewrite"]
    deletes  = [op for op in operations if op.get("op") == "delete"]
    inserts  = [op for op in operations if op.get("op") == "insert_after"]

    # Rewrites
    for op in rewrites:
        ref = op.get("ref", "")
        idx = resolve(ref)
        elem = get_para_elem(idx)
        if elem is None:
            errors.append(f"❌ {ref} не найден (idx={idx}, total={len(all_paras_list)})")
            continue
        old = "".join(t.text or "" for t in elem.iter(_qn("w:t")))[:60]
        new_text = op.get("new_text", "")
        runs = list(elem.iter(_qn("w:r")))
        if runs:
            # Set text in first run, clear the rest
            t_elems = list(runs[0].iter(_qn("w:t")))
            if t_elems:
                t_elems[0].text = new_text
                for t in t_elems[1:]:
                    t.text = ""
            else:
                t_new = etree.SubElement(runs[0], _qn("w:t"))
                t_new.text = new_text
            for r in runs[1:]:
                for t in r.iter(_qn("w:t")):
                    t.text = ""
        done.append(f"✏️ {ref}: «{old}» → «{new_text[:60]}»")

    # Deletes (in reverse index order to keep positions stable)
    delete_items = []
    for op in deletes:
        ref = op.get("ref", "")
        idx = resolve(ref)
        elem = get_para_elem(idx)
        if elem is None:
            errors.append(f"❌ {ref} не найден")
            continue
        old = "".join(t.text or "" for t in elem.iter(_qn("w:t")))[:60]
        delete_items.append((idx, ref, old, elem))

    for idx, ref, old, elem in sorted(delete_items, key=lambda x: x[0], reverse=True):
        if is_in_table(elem):
            # OOXML: w:tc must have at least one w:p — just clear the text
            for t in elem.iter(_qn("w:t")):
                t.text = ""
            done.append(f"🗑️ {ref} (в таблице — очищен): «{old}»")
        else:
            parent = elem.getparent()
            if parent is not None:
                parent.remove(elem)
                done.append(f"🗑️ {ref}: «{old}»")
            else:
                errors.append(f"❌ {ref}: не удалось удалить")

    # Inserts (new paragraph after anchor or after containing table)
    for op in inserts:
        ref = op.get("ref", "")
        idx = resolve(ref)
        anchor = get_para_elem(idx)
        if anchor is None:
            errors.append(f"❌ {ref} не найден")
            continue
        text = op.get("text", "")

        # If anchor is inside a table, insert after the whole table to avoid corrupting it
        if is_in_table(anchor):
            insert_after_elem = get_containing_table(anchor)
        else:
            insert_after_elem = anchor

        # Build new paragraph copying style from a top-level paragraph (first body-level w:p)
        body_paras = [p for p in all_paras_list if not is_in_table(p)]
        style_source = body_paras[0] if body_paras else anchor
        new_p = copy.deepcopy(style_source)
        for t in new_p.iter(_qn("w:t")):
            t.text = ""
        runs = list(new_p.iter(_qn("w:r")))
        if runs:
            t_elem = runs[0].find(_qn("w:t"))
            if t_elem is None:
                t_elem = etree.SubElement(runs[0], _qn("w:t"))
            t_elem.text = text
            for r in runs[1:]:
                p = r.getparent()
                if p is not None:
                    p.remove(r)
        insert_after_elem.addnext(new_p)
        location = "после таблицы" if is_in_table(anchor) else f"после {ref}"
        done.append(f"➕ {location}: «{text[:60]}»")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue(), _build_diff(done, errors)


# ── python-docx fallback ───────────────────────────────────────────────────

def _para_hash(text: str) -> str:
    import hashlib
    return hashlib.md5(text.encode()).hexdigest()[:4]


def _read_fallback(file_bytes: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    lines = []
    idx = 1

    for para in doc.paragraphs:
        text = para.text.strip()
        h = _para_hash(text or f"__empty_{idx}__")
        preview = text[:100] + ("..." if len(text) > 100 else "")
        lines.append(f"P{idx}#{h}| {preview}")
        idx += 1

    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    text = para.text.strip()
                    if text:
                        h = _para_hash(text)
                        preview = text[:80] + ("..." if len(text) > 80 else "")
                        lines.append(f"T{r_idx+1}C{c_idx+1}#{h}| {preview}")

    if not lines:
        return "Документ пустой."

    return f"Всего параграфов: {len(lines)}\n\n" + "\n".join(lines)


def _apply_fallback(file_bytes: bytes, operations: list[dict], filename: str) -> tuple[bytes, str]:
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    done = []
    errors = []

    # Build hash→paragraph index map
    para_map: dict[str, int] = {}
    for i, para in enumerate(doc.paragraphs):
        h = _para_hash(para.text.strip() or f"__empty_{i}__")
        key = f"P{i+1}#{h}"
        para_map[key] = i

    def find_para_idx(ref: str) -> int:
        """Find paragraph index by ref, tolerating index mismatch (hash-only match)."""
        if ref in para_map:
            return para_map[ref]
        # Hash-only fallback: match any ref with same hash suffix
        hash_suffix = ref.split("#")[1] if "#" in ref else ""
        for k, v in para_map.items():
            if k.endswith(f"#{hash_suffix}"):
                return v
        return -1

    for op in operations:
        kind = op.get("op", "rewrite")
        ref = op.get("ref", "")
        idx = find_para_idx(ref)

        try:
            if kind == "rewrite":
                if idx < 0:
                    errors.append(f"❌ {ref} не найден")
                    continue
                para = doc.paragraphs[idx]
                new_text = op.get("new_text", "")
                old_text = para.text[:60]
                if para.runs:
                    para.runs[0].text = new_text
                    for r in para.runs[1:]:
                        r.text = ""
                done.append(f"✏️ {ref}: «{old_text}» → «{new_text[:60]}»")

            elif kind == "delete":
                if idx < 0:
                    errors.append(f"❌ {ref} не найден")
                    continue
                para = doc.paragraphs[idx]
                old_text = para.text[:60]
                p = para._element
                p.getparent().remove(p)
                done.append(f"🗑️ {ref}: удалён «{old_text}»")

            elif kind == "insert_after":
                if idx < 0:
                    errors.append(f"❌ {ref} не найден")
                    continue
                import copy
                from lxml import etree
                from docx.oxml.ns import qn
                text = op.get("text", "")
                anchor = doc.paragraphs[idx]._element
                new_p = copy.deepcopy(anchor)
                for t in new_p.findall(".//" + qn("w:t")):
                    t.text = ""
                runs = new_p.findall(".//" + qn("w:r"))
                if runs:
                    t_elem = runs[0].find(qn("w:t"))
                    if t_elem is None:
                        t_elem = etree.SubElement(runs[0], qn("w:t"))
                    t_elem.text = text
                anchor.addnext(new_p)
                done.append(f"➕ После {ref}: вставлено «{text[:60]}»")

        except Exception as e:
            errors.append(f"❌ {ref}: {e}")
            logger.error("Fallback op %s on %s failed: %s", kind, ref, e)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue(), _build_diff(done, errors)


def _build_diff(done: list[str], errors: list[str]) -> str:
    parts = []
    if done:
        parts.append(f"Выполнено ({len(done)}):\n" + "\n".join(done))
    if errors:
        parts.append(f"Ошибки ({len(errors)}):\n" + "\n".join(errors))
    return "\n\n".join(parts) or "Операций не выполнено."
